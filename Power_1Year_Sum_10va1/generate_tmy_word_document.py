"""
生成 TMY 技術文件的 Word 版本
執行方式: python generate_tmy_word_document.py
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from datetime import datetime

def set_cell_background(cell, fill_color):
    """設定表格儲存格背景顏色"""
    cell_properties = cell._element.get_or_add_tcPr()
    shade_obj = OxmlElement('w:shd')
    shade_obj.set(qn('w:fill'), fill_color)
    cell_properties.append(shade_obj)

def add_heading_with_line(doc, text, level=1):
    """添加標題"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

def create_tmy_word_document():
    """建立完整的 TMY 技術文件"""
    
    doc = Document()
    
    # 設定預設字型 (支援中文)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # ===== 標題頁 =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('典型氣象年 (TMY) 資料在未來外生變數中的處理方法')
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    
    # 空行
    doc.add_paragraph()
    
    # 文件資訊
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    info_data = [
        ('文件版本', '1.0'),
        ('撰寫日期', '2026 年 5 月 1 日'),
        ('應用場景', 'AutoTS 模型的 365 天未來日發電量預測'),
        ('目標讀者', '學位論文、技術報告、同行評審'),
    ]
    
    for i, (key, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = key
        info_table.rows[i].cells[1].text = value
        set_cell_background(info_table.rows[i].cells[0], 'D3D3D3')
    
    doc.add_paragraph()
    
    # ===== 1. 概述 =====
    add_heading_with_line(doc, '1. 概述', level=1)
    
    add_heading_with_line(doc, '1.1 動機與目標', level=2)
    doc.add_paragraph(
        '在進行長期 (365 天) 時間序列預測時，模型需要未來的外生變數 (exogenous variables) '
        '以提供充分的特徵資訊。由於無法獲得真實未來氣象觀測值，本研究採用典型氣象年 '
        '(Typical Meteorological Year, TMY) 資料作為未來期間的外生變數代理。'
    )
    
    add_heading_with_line(doc, '1.2 核心概念', level=2)
    doc.add_paragraph(
        '典型氣象年 (TMY) 是一組代表特定位置多年氣象特徵的日尺度資料集。其構成方法是：'
    )
    doc.add_paragraph('從多年氣象觀測資料 (通常 15-20 年) 中，篩選具有代表性的典型日子', style='List Bullet')
    doc.add_paragraph('將這些典型日子拼接成一個完整年份 (365 天)', style='List Bullet')
    doc.add_paragraph('每個月份的每一天都選取多年觀測資料中「最具代表性」的值', style='List Bullet')
    
    doc.add_paragraph('關鍵特性：')
    doc.add_paragraph('平滑性：消除單一年份的極端事件和異常值', style='List Bullet')
    doc.add_paragraph('統計代表性：保留多年平均的月季節特徵', style='List Bullet')
    doc.add_paragraph('無未來信息：完全基於歷史觀測，不含預測成分', style='List Bullet')
    
    # ===== 2. 資料來源與規格 =====
    add_heading_with_line(doc, '2. 資料來源與規格', level=1)
    
    add_heading_with_line(doc, '2.1 TMY 資料集特徵', level=2)
    
    tmy_table = doc.add_table(rows=7, cols=2)
    tmy_table.style = 'Light Grid Accent 1'
    
    tmy_data = [
        ('特徵項目', '值'),
        ('地點座標', '24.148°N, 120.703°E (台北，台灣)'),
        ('資料時間範圍', '2005-2023 (19 年歷史觀測)'),
        ('時間尺度', '日尺度 (Daily)'),
        ('時區', 'UTC+8 (台灣標準時間)'),
        ('記錄筆數', '365 日/年 × 1 個 TMY 年份 = 365 筆'),
        ('資料來源', 'PVGIS (Photovoltaic Geographical Information System)'),
    ]
    
    for i, (key, value) in enumerate(tmy_data):
        tmy_table.rows[i].cells[0].text = key
        tmy_table.rows[i].cells[1].text = value
        if i == 0:
            set_cell_background(tmy_table.rows[i].cells[0], 'D3D3D3')
            set_cell_background(tmy_table.rows[i].cells[1], 'D3D3D3')
    
    add_heading_with_line(doc, '2.2 檔案規格', level=2)
    doc.add_paragraph('檔案名稱：')
    code_para = doc.add_paragraph('tmy_24.148_120.703_2005_2023[UTC+8][daily][mapped][dateAdj].csv')
    code_para.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_paragraph('主要欄位 (本研究使用的部分)：')
    doc.add_paragraph('date：日期 (格式: YYYY/MM/DD)', style='List Bullet')
    doc.add_paragraph('Temperature：日均溫度 (°C)', style='List Bullet')
    doc.add_paragraph('RH：日均相對濕度 (%)', style='List Bullet')
    
    # ===== 3. 外生變數選擇 =====
    add_heading_with_line(doc, '3. 外生變數選擇的科學依據', level=1)
    
    add_heading_with_line(doc, '3.1 為什麼選擇溫度和相對濕度', level=2)
    
    doc.add_paragraph('溫度 (Temperature, T) 對光伏發電的影響：')
    doc.add_paragraph('光伏電池效率與工作溫度呈負相關', style='List Bullet')
    doc.add_paragraph('台灣夏季高溫 (30-35°C) 時，發電效率下降 15-25%', style='List Bullet')
    doc.add_paragraph('冬季低溫 (10-15°C) 時，效率相對最高', style='List Bullet')
    
    doc.add_paragraph('相對濕度 (RH) 對光伏發電的影響：')
    doc.add_paragraph('高濕度增加大氣淵度，削弱輻照度', style='List Bullet')
    doc.add_paragraph('台灣夏季 RH 常在 70-80%，冬季 50-70%', style='List Bullet')
    doc.add_paragraph('霧霾與高濕度可能同時出現，進一步衰減輻照度', style='List Bullet')
    
    # ===== 4. 資料預處理 =====
    add_heading_with_line(doc, '4. 資料預處理流程', level=1)
    
    add_heading_with_line(doc, '4.1 處理步驟', level=2)
    
    doc.add_paragraph('步驟 1: 檔案讀取與欄位偵測')
    code1 = doc.add_paragraph(
        'import pandas as pd\n'
        'from pathlib import Path\n\n'
        'CSV_PATH = Path("tmy_24.148_120.703_2005_2023[UTC+8][daily][mapped][dateAdj].csv")\n'
        'df = pd.read_csv(CSV_PATH)\n'
        'date_col = next((col for col in ("date", "Date", "LocalTime") if col in df.columns), None)'
    )
    code1.style = 'No Spacing'
    code1_format = code1.paragraph_format
    code1_format.left_indent = Inches(0.5)
    code1_format.space_before = Pt(6)
    code1_format.space_after = Pt(6)
    for run in code1.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 0, 0)
    
    doc.add_paragraph(
        '目的：確保檔案格式與預期一致，自動識別日期欄位名稱'
    )
    
    doc.add_paragraph()
    doc.add_paragraph('步驟 2: 時間轉換與標準化')
    code2 = doc.add_paragraph(
        'df[date_col] = pd.to_datetime(df[date_col], errors="coerce")\n'
        'REGRESSOR_COLUMNS = ["Temperature", "RH"]\n'
        'out = df[[date_col] + REGRESSOR_COLUMNS].copy()\n'
        'for col in REGRESSOR_COLUMNS:\n'
        '    out[col] = pd.to_numeric(out[col], errors="coerce")'
    )
    code2.style = 'No Spacing'
    code2_format = code2.paragraph_format
    code2_format.left_indent = Inches(0.5)
    for run in code2.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 0, 0)
    
    doc.add_paragraph('預期行為：無法轉換的值變成 NaN (便於後續檢查)')
    
    doc.add_paragraph()
    doc.add_paragraph('步驟 3: 去重與排序')
    code3 = doc.add_paragraph(
        'out = out.dropna(subset=[date_col]).sort_values(date_col)\n'
        'out = out.drop_duplicates(subset=[date_col], keep="last")\n'
        'out = out.rename(columns={date_col: "date"}).set_index("date")\n'
        'out = out.asfreq("D")'
    )
    code3.style = 'No Spacing'
    code3_format = code3.paragraph_format
    code3_format.left_indent = Inches(0.5)
    for run in code3.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 0, 0)
    
    doc.add_paragraph('預期行為：TMY 資料恰好包含 365 筆記錄 (non-leap year)')
    
    # ===== 5. 未來期間外生變數構造 =====
    add_heading_with_line(doc, '5. 未來期間外生變數構造', level=1)
    
    add_heading_with_line(doc, '5.1 問題定義', level=2)
    doc.add_paragraph('輸入：TMY 資料 (一個完整年份的典型日) + 預測時間索引 (未來 365 日)', style='List Bullet')
    doc.add_paragraph('輸出：未來外生變數矩陣 (365, 2)，欄位為 [Temperature, RH]', style='List Bullet')
    doc.add_paragraph('核心挑戰：直接重複 TMY 資料無法適應未來的月份序列', style='List Bullet')
    
    add_heading_with_line(doc, '5.2 月-日氣候平均值策略', level=2)
    
    doc.add_paragraph('核心想法：將 TMY 視為月份-日子的函數，不依賴年份資訊')
    
    doc.add_paragraph('具體步驟：')
    
    doc.add_paragraph('對 TMY 資料按月-日 (month-day) 分組', style='List Number')
    code_step1 = doc.add_paragraph(
        'tmy_tmp = tmy_df[REGRESSOR_COLUMNS].copy()\n'
        'tmy_tmp["md"] = tmy_tmp.index.strftime("%m-%d")\n'
        'by_md = tmy_tmp.groupby("md")[REGRESSOR_COLUMNS].mean()'
    )
    code_step1.style = 'No Spacing'
    code_step1_format = code_step1.paragraph_format
    code_step1_format.left_indent = Inches(0.5)
    for run in code_step1.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 0, 0)
    
    doc.add_paragraph('結果：365 列的表格，每列對應一個月-日，包含平均溫度與平均濕度')
    
    doc.add_paragraph('將未來日期重新索引', style='List Number')
    code_step2 = doc.add_paragraph(
        'forecast_index = pd.date_range(\n'
        '    start="2026-03-01",\n'
        '    periods=365,\n'
        '    freq="D"\n'
        ')\n'
        'md_keys = forecast_index.strftime("%m-%d")'
    )
    code_step2.style = 'No Spacing'
    code_step2_format = code_step2.paragraph_format
    code_step2_format.left_indent = Inches(0.5)
    for run in code_step2.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 0, 0)
    
    doc.add_paragraph('查表填值', style='List Number')
    code_step3 = doc.add_paragraph(
        'fill_by_md = by_md.reindex(md_keys)\n'
        'fill_by_md.index = forecast_index'
    )
    code_step3.style = 'No Spacing'
    code_step3_format = code_step3.paragraph_format
    code_step3_format.left_indent = Inches(0.5)
    for run in code_step3.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 0, 0)
    
    # ===== 6. 與其他方法比較 =====
    add_heading_with_line(doc, '6. 與其他方法的比較', level=1)
    
    comparison_table = doc.add_table(rows=6, cols=5)
    comparison_table.style = 'Light Grid Accent 1'
    
    comparison_data = [
        ('方法', '實現難度', '計算成本', '準確性', '科學合理性'),
        ('月-日平均 (本方法)', '低', '極低', '中', '★★★★☆'),
        ('年份重複', '極低', '極低', '低', '★★☆☆☆'),
        ('氣象預報模型', '高', '高', '高', '★★★★★'),
        ('機器學習 (LSTM)', '高', '中-高', '中-高', '★★★☆☆'),
    ]
    
    for i, row_data in enumerate(comparison_data):
        for j, cell_data in enumerate(row_data):
            comparison_table.rows[i].cells[j].text = cell_data
            if i == 0:
                set_cell_background(comparison_table.rows[i].cells[j], 'D3D3D3')
    
    # ===== 7. 驗證與品質保證 =====
    add_heading_with_line(doc, '7. 驗證與品質保證', level=1)
    
    add_heading_with_line(doc, '7.1 資料完整性檢查', level=2)
    
    verify_code = doc.add_paragraph(
        '# 檢查 1: 記錄數\n'
        'assert len(tmy_df) in [365, 366], "Record count must be 365 or 366"\n\n'
        '# 檢查 2: 無缺值\n'
        'assert tmy_df.isna().sum().sum() == 0, "TMY data should have no missing values"\n\n'
        '# 檢查 3: 溫度在合理範圍\n'
        'assert (tmy_df["Temperature"] > -20).all() and (tmy_df["Temperature"] < 50).all()\n\n'
        '# 檢查 4: 濕度在 0-100% 範圍\n'
        'assert (tmy_df["RH"] >= 0).all() and (tmy_df["RH"] <= 100).all()'
    )
    verify_code.style = 'No Spacing'
    verify_code_format = verify_code.paragraph_format
    verify_code_format.left_indent = Inches(0.5)
    for run in verify_code.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 0, 0)
    
    # ===== 8. 最佳實踐 =====
    add_heading_with_line(doc, '8. 最佳實踐', level=1)
    
    add_heading_with_line(doc, '8.1 程式碼結構建議', level=2)
    
    best_code = doc.add_paragraph(
        'class TMYProcessor:\n'
        '    """TMY 資料處理器。"""\n'
        '    def __init__(self, tmy_path: Path, logger):\n'
        '        self.tmy_path = tmy_path\n'
        '        self.tmy_df = None\n'
        '    def load(self) -> pd.DataFrame:\n'
        '        self.tmy_df = read_tmy_data(self.tmy_path)\n'
        '        return self.tmy_df\n'
        '    def build_future_regressor(self, forecast_index):\n'
        '        return build_predict_regressor(self.tmy_df, forecast_index)'
    )
    best_code.style = 'No Spacing'
    best_code_format = best_code.paragraph_format
    best_code_format.left_indent = Inches(0.5)
    for run in best_code.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 0, 0)
    
    # ===== 9. 結論 =====
    add_heading_with_line(doc, '9. 結論', level=1)
    
    add_heading_with_line(doc, '9.1 技術總結', level=2)
    doc.add_paragraph('本文件說明如何利用 TMY 資料為 365 天預測提供外生變數：')
    doc.add_paragraph('資料選擇：溫度 (T) 和相對濕度 (RH)', style='List Number')
    doc.add_paragraph('預處理：時間轉換、去重、頻率設定', style='List Number')
    doc.add_paragraph('未來映射：月-日氣候平均值法', style='List Number')
    doc.add_paragraph('品質驗證：多層次驗證', style='List Number')
    
    add_heading_with_line(doc, '9.2 適用條件', level=2)
    doc.add_paragraph('適用於：')
    doc.add_paragraph('中期預測 (1-12 個月)', style='List Bullet')
    doc.add_paragraph('以季節性為主的場景', style='List Bullet')
    doc.add_paragraph('計算資源有限的環境', style='List Bullet')
    
    doc.add_paragraph('不適用於：')
    doc.add_paragraph('超長期預測 (>5 年)', style='List Bullet')
    doc.add_paragraph('特定年份特異事件', style='List Bullet')
    doc.add_paragraph('日內 (小時級) 精度', style='List Bullet')
    
    add_heading_with_line(doc, '9.3 改進方向', level=2)
    doc.add_paragraph('適應性月-日平均 (依預測年份調整權重)', style='List Number')
    doc.add_paragraph('多模式集合 (多個 TMY 資料集)', style='List Number')
    doc.add_paragraph('機器學習增強 (訓練修正模型)', style='List Number')
    doc.add_paragraph('邊界修正 (leap day 特殊處理)', style='List Number')
    
    # ===== 頁尾 =====
    doc.add_paragraph()
    footer_para = doc.add_paragraph('文件結束')
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    final_para = doc.add_paragraph(
        f'最後更新: {datetime.now().strftime("%Y 年 %m 月 %d 日")}\n'
        '應用場景: AutoTS 365 天日發電量預測 (台北)'
    )
    final_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    final_para_format = final_para.paragraph_format
    final_para_format.space_before = Pt(12)
    
    # 保存檔案
    output_path = Path('TMY_ExogenousVariableProcessing_Technical_Document.docx')
    doc.save(output_path)
    print(f'✓ Word 檔案已生成: {output_path}')
    print(f'✓ 檔案大小: {output_path.stat().st_size / 1024:.1f} KB')
    return output_path

if __name__ == '__main__':
    try:
        create_tmy_word_document()
    except ImportError as e:
        print(f'Error: {e}')
        print('Please install python-docx: pip install python-docx')
    except Exception as e:
        print(f'Error: {e}')
